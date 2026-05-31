"""
文档解析与文本切片服务

功能：
  - 从 PDF/Word/TXT/Markdown 文件中提取纯文本
  - 将长文本按段落和句子边界切分为多个小片段（chunk）

这是 RAG（检索增强生成）系统的核心步骤：
  1. 文档上传 → 提取文本 → 切分为小片段
  2. 每个小片段通过 Embedding 模型转化为向量
  3. 用户提问时，检索最相关的片段，交给大模型生成回答
"""

import re

import fitz  # PyMuPDF — 用于解析 PDF 文件（需安装：pip install PyMuPDF）
from docx import (
    Document as DocxDocument,  # python-docx — 用于解析 .docx 文件（需安装：pip install python-docx）
)


def extract_text_with_pages(filepath: str, file_type: str) -> list[tuple[str, int]]:
    """
    提取文本并追踪每段文本的页码。

    返回：
        [(文本段落, 页码), ...] 列表。页码从 1 开始。
        非 PDF 文件所有文本归为第 1 页。
    """
    if file_type == "pdf":
        return _extract_pdf_with_pages(filepath)
    elif file_type == "docx":
        text = _extract_docx(filepath)
        return [(text, 1)] if text.strip() else []
    elif file_type in ("txt", "md"):
        text = _extract_text_file(filepath)
        return [(text, 1)] if text.strip() else []
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def _extract_pdf_with_pages(filepath: str) -> list[tuple[str, int]]:
    """提取 PDF 文本，保留页码信息。"""
    doc = fitz.open(filepath)
    try:
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                pages.append((text, i + 1))
        return pages
    finally:
        doc.close()


def extract_text(filepath: str, file_type: str) -> str:
    """
    根据文件类型提取文本内容。

    参数：
        filepath: 文件的完整路径
        file_type: 文件类型，支持 pdf / docx / txt / md

    返回：
        提取出的纯文本字符串
    """
    if file_type == "pdf":
        return _extract_pdf(filepath)
    elif file_type == "docx":
        return _extract_docx(filepath)
    elif file_type in ("txt", "md"):
        return _extract_text_file(filepath)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def _extract_pdf(filepath: str) -> str:
    """
    提取 PDF 文件中的文本。

    PyMuPDF (fitz) 的工作机制：
      - fitz.open() 打开 PDF 文件，返回一个 Document 对象
        （底层解析 PDF 的二进制结构，提取文本、图片、字体等信息）
      - Document 对象可按页迭代，每一页是一个 Page 对象
      - page.get_text() 从当前页面中提取所有可见文本
        （它会按照 PDF 内部的文本绘制顺序来拼接文字，因此可能偶尔出现顺序问题）
      - 最后将所有页面的文本用换行符连接，保留原始的页面分隔
    """
    doc = fitz.open(filepath)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _extract_docx(filepath: str) -> str:
    """
    提取 Word (.docx) 文件中的文本。

    python-docx 的工作机制：
      - .docx 文件本质上是一个 ZIP 压缩包，内部包含多个 XML 文件
        （你可以将 .docx 后缀改为 .zip 后用解压工具查看其内部结构）
      - DocxDocument() 解析这个 ZIP 包，读取 word/document.xml 中的段落和表格信息
      - document.paragraphs 是一个列表，包含文档中的所有段落对象
      - 每个段落的 .text 属性就是该段落的纯文本内容
      - 使用 strip() 过滤掉空段落（有些段落仅用于排版，不含实际文字）
    """
    doc = DocxDocument(filepath)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text_file(filepath: str) -> str:
    """
    读取纯文本文件（TXT 或 Markdown）。

    这是最简单的情况：直接按 UTF-8 编码读取文件内容即可。
    使用 UTF-8 编码是因为它兼容 ASCII，且能正确处理中文等多字节字符。
    如果遇到编码问题，可以考虑改用 chardet 库自动检测编码。
    """
    with open(filepath, encoding="utf-8") as f:
        return f.read()


# 中英文句子结束标点
_SENTENCE_ENDINGS = re.compile(r"(?<=[。！？.!?；\n])")


def _split_by_sentence(paragraph: str, chunk_size: int) -> list[str]:
    """将单个段落按句子边界切割，超长单句按字符兜底。"""
    # 按句子标点切分，保留标点
    sentences = [s.strip() for s in _SENTENCE_ENDINGS.split(paragraph) if s.strip()]

    # 没有句子标点，直接按字符兜底
    if not sentences:
        return [paragraph[i : i + chunk_size] for i in range(0, len(paragraph), chunk_size)] or []

    chunks = []
    current = ""
    for sent in sentences:
        if not current:
            current = sent
        elif len(current) + len(sent) <= chunk_size:
            current += sent
        else:
            if current:
                chunks.append(current)
            # 单句超长，按字符兜底切割
            if len(sent) > chunk_size:
                for i in range(0, len(sent), chunk_size):
                    piece = sent[i : i + chunk_size]
                    if piece:
                        chunks.append(piece)
                current = ""
            else:
                current = sent
    if current:
        chunks.append(current)
    return chunks


def split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    将长文本切分为多个小片段（chunk），优先在段落和句子边界处切割。

    切割策略（由粗到细）：
      1. 先按段落分（双换行 \\n\\n）
      2. 如果某段超过 chunk_size，按句子分（。！？.!?）
      3. 如果单句仍超长，按字符数兜底
      4. 相邻 chunk 之间保留 overlap（取前一个 chunk 的末尾句子）

    参数：
        text: 原始文本
        chunk_size: 每个片段的最大字符数（默认 500）
        overlap: 相邻片段之间的重叠字符数（默认 50）

    返回：
        切分后的片段列表，每个元素是一个非空字符串
    """
    if not text.strip():
        return []

    if overlap >= chunk_size:
        overlap = chunk_size - 1

    # 第一步：按段落分（双换行为段落分隔）
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

    # 第二步：处理每个段落，超长段落按句子切割
    # 每个段落（或其切分后的子片段）成为一个独立 chunk，不跨段落合并
    chunks = []
    for para in paragraphs:
        if len(para) <= chunk_size:
            chunks.append(para)
        else:
            sub_chunks = _split_by_sentence(para, chunk_size)
            for sc in sub_chunks:
                if len(sc) <= chunk_size:
                    chunks.append(sc)
                else:
                    # 句子切割后仍超长，按字符兜底
                    for i in range(0, len(sc), chunk_size):
                        piece = sc[i : i + chunk_size]
                        if piece:
                            chunks.append(piece)

    # 第三步：添加 overlap（取前一个 chunk 的末尾句子作为下一个 chunk 的开头）
    if len(chunks) <= 1:
        return chunks

    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev = chunks[i - 1]
        curr = chunks[i]
        # 从上一个 chunk 尾部提取末尾句子作为 overlap
        overlap_text = _extract_tail(prev, overlap)
        if overlap_text:
            combined = overlap_text + curr
            # 如果合并后超长，只保留能放下的部分
            if len(combined) <= chunk_size:
                result.append(combined)
            else:
                result.append(curr)
        else:
            result.append(curr)
    return result


def _extract_tail(text: str, max_len: int) -> str:
    """提取文本末尾的内容作为 overlap，优先取完整句子。"""
    if len(text) <= max_len:
        return text
    tail = text[-max_len:]
    # 尝试从句子边界开始
    for ending in ("。", "！", "？", ".", "!", "?", "；", "\n"):
        idx = tail.find(ending)
        if idx >= 0:
            return tail[idx + 1 :]
    return tail


# 标题检测正则
_HEADER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百千万\d]+[章节篇部条]"),  # 第一章、第二节
    re.compile(r"^\d+(?:\.\d+)*[.\s]"),  # 1. / 2.3 / 3.1.2
    re.compile(r"^#{1,6}\s"),  # Markdown 标题
]


def _is_header_line(line: str) -> bool:
    """判断一行文本是否是标题。"""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    return any(pat.match(stripped) for pat in _HEADER_PATTERNS)


def _detect_headers(text: str) -> list[tuple[int, str]]:
    """
    检测文本中所有标题及其字符偏移位置。

    返回：
        [(偏移位置, 标题文本), ...]
    """
    headers = []
    for match in re.finditer(r"[^\n]+", text):
        line = match.group()
        if _is_header_line(line):
            headers.append((match.start(), line.strip()))
    return headers


def _build_title_path(text: str, char_offset: int, headers: list[tuple[int, str]]) -> str:
    """
    根据字符偏移位置，构建当前位置的标题层级路径。

    参数：
        text: 完整文本
        char_offset: 当前 chunk 在原文中的字符偏移
        headers: _detect_headers 返回的 [(偏移, 标题文本), ...]

    返回：
        标题层级路径，如 "第二章 薪酬 > 2.1 基本工资"
    """
    if not headers:
        return ""

    # 找到 offset 之前（含）的所有标题
    applicable = [(off, title) for off, title in headers if off <= char_offset]
    if not applicable:
        return ""

    return " > ".join(title for _, title in applicable)


def _split_by_structure(text: str, chunk_size: int = 500, overlap: int = 50) -> list[dict]:
    """
    结构感知递归切分：按标题→段落→句子逐级切割。

    返回：
        [{"text": str, "title_path": str, "char_offset": int}, ...]
    """
    headers = _detect_headers(text)

    if not headers:
        # 无标题，退化为普通切分
        plain_chunks = split_text(text, chunk_size=chunk_size, overlap=overlap)
        result = []
        offset = 0
        for chunk in plain_chunks:
            idx = text.find(chunk, offset)
            actual_offset = idx if idx >= 0 else offset
            result.append(
                {
                    "text": chunk,
                    "title_path": "",
                    "char_offset": actual_offset,
                }
            )
            offset = actual_offset + len(chunk)
        return result

    # 按标题分割文本为 sections
    sections = []
    for i, (offset, title) in enumerate(headers):
        end = headers[i + 1][0] if i + 1 < len(headers) else len(text)
        section_text = text[offset:end].strip()
        if section_text:
            sections.append((offset, title, section_text))

    # 处理标题之前的内容（前言）
    first_header_offset = headers[0][0]
    if first_header_offset > 0:
        preamble = text[:first_header_offset].strip()
        if preamble:
            sections.insert(0, (0, "", preamble))

    # 对每个 section 进行段落/句子级切分
    # 结构边界已提供语义分隔，section 内部不需要 overlap（避免偏移追踪失准）
    result = []
    for section_offset, _title, section_text in sections:
        sub_chunks = split_text(section_text, chunk_size=chunk_size, overlap=0)
        for chunk in sub_chunks:
            idx = section_text.find(chunk)
            actual_offset = section_offset + (idx if idx >= 0 else 0)
            title_path = _build_title_path(text, actual_offset, headers)
            result.append(
                {
                    "text": chunk,
                    "title_path": title_path,
                    "char_offset": actual_offset,
                }
            )

    return result


def _split_by_semantics(
    text: str, chunk_size: int = 500, threshold_percentile: float = 25
) -> list[dict]:
    """
    语义切分：基于相邻句子的 embedding 相似度骤降处切分。

    参数：
        text: 待切分文本
        chunk_size: 单个 chunk 最大字符数
        threshold_percentile: 相似度阈值百分位，低于此百分位的相似度处切分

    返回：
        [{"text": str, "title_path": "", "char_offset": int}, ...]
    """
    if not text.strip():
        return []

    # 按句子切分
    sentences = [s.strip() for s in _SENTENCE_ENDINGS.split(text) if s.strip()]
    if len(sentences) <= 1:
        return [{"text": text.strip(), "title_path": "", "char_offset": 0}]

    # 延迟导入：仅语义切分需要 numpy 和 embedding 服务
    import numpy as np

    from app.services.embedding import get_embeddings

    # 计算每个句子的 embedding
    embeddings = get_embeddings(sentences)
    embeddings_np = np.array(embeddings)

    # 计算相邻句子的余弦相似度
    similarities = []
    for i in range(len(embeddings_np) - 1):
        a, b = embeddings_np[i], embeddings_np[i + 1]
        dot = np.dot(a, b)
        norm = np.linalg.norm(a) * np.linalg.norm(b)
        sim = dot / norm if norm > 0 else 0.0
        similarities.append(sim)

    if not similarities:
        return [{"text": " ".join(sentences), "title_path": "", "char_offset": 0}]

    # 在相似度低于阈值处切分
    threshold = np.percentile(similarities, threshold_percentile)
    split_indices = [i for i, sim in enumerate(similarities) if sim < threshold]

    # 按切分点分组句子
    groups = []
    prev = 0
    for idx in split_indices:
        groups.append(sentences[prev : idx + 1])
        prev = idx + 1
    if prev < len(sentences):
        groups.append(sentences[prev:])

    # 合并过短的组，拆分过长的组
    result = []
    current_offset = 0
    for group in groups:
        chunk_text = "".join(group)
        # 如果超长，用普通切分兜底
        if len(chunk_text) > chunk_size:
            sub_chunks = split_text(chunk_text, chunk_size=chunk_size)
            for sc in sub_chunks:
                idx = text.find(sc, current_offset)
                actual_offset = idx if idx >= 0 else current_offset
                result.append({"text": sc, "title_path": "", "char_offset": actual_offset})
                current_offset = actual_offset + len(sc)
        else:
            idx = text.find(chunk_text, current_offset)
            actual_offset = idx if idx >= 0 else current_offset
            result.append({"text": chunk_text, "title_path": "", "char_offset": actual_offset})
            current_offset = actual_offset + len(chunk_text)

    return result


# 混合策略阈值：低于此字符数的文档使用简单切分
_SHORT_DOC_THRESHOLD = 1000


def split_document(
    text: str,
    file_type: str = "txt",
    chunk_size: int = 500,
    overlap: int = 50,
    pages: list[tuple[str, int]] | None = None,
) -> list[dict]:
    """
    统一文档切分入口，根据文档特征自动选择切分策略。

    策略选择逻辑：
      1. 有标题结构 → 结构感知切分
      2. 无标题 + 短文档（< 1000字）→ 普通段落/句子切分
      3. 无标题 + 长文档 → 语义切分

    参数：
        text: 完整文本（当传入 pages 时可为空字符串）
        file_type: 文件类型
        chunk_size: chunk 最大字符数
        overlap: 重叠字符数
        pages: PDF 页面数据 [(页文本, 页码), ...]，为 None 时从 text 切分

    返回：
        [{"text": str, "title_path": str, "char_offset": int, "page": int}, ...]
    """
    if pages:
        return _split_pages(pages, chunk_size=chunk_size, overlap=overlap)

    if not text.strip():
        return []

    headers = _detect_headers(text)

    if headers:
        chunks = _split_by_structure(text, chunk_size=chunk_size, overlap=overlap)
        for chunk in chunks:
            chunk["page"] = 1
        return chunks

    if len(text) < _SHORT_DOC_THRESHOLD:
        plain_chunks = split_text(text, chunk_size=chunk_size, overlap=overlap)
        result = []
        offset = 0
        for chunk in plain_chunks:
            idx = text.find(chunk, offset)
            actual_offset = idx if idx >= 0 else offset
            result.append(
                {
                    "text": chunk,
                    "title_path": "",
                    "char_offset": actual_offset,
                    "page": 1,
                }
            )
            offset = actual_offset + len(chunk)
        return result

    # 长文档无标题 → 语义切分
    chunks = _split_by_semantics(text, chunk_size=chunk_size)
    for chunk in chunks:
        chunk["page"] = 1
    return chunks


def _split_pages(
    pages: list[tuple[str, int]], chunk_size: int = 500, overlap: int = 50
) -> list[dict]:
    """按页面分割，每页独立走切分策略。"""
    result = []
    for page_text, page_num in pages:
        if not page_text.strip():
            continue
        headers = _detect_headers(page_text)
        if headers:
            chunks = _split_by_structure(page_text, chunk_size=chunk_size, overlap=overlap)
        elif len(page_text) < _SHORT_DOC_THRESHOLD:
            plain = split_text(page_text, chunk_size=chunk_size, overlap=overlap)
            chunks = []
            offset = 0
            for c in plain:
                idx = page_text.find(c, offset)
                actual_offset = idx if idx >= 0 else offset
                chunks.append({"text": c, "title_path": "", "char_offset": actual_offset})
                offset = actual_offset + len(c)
        else:
            chunks = _split_by_semantics(page_text, chunk_size=chunk_size)
        for chunk in chunks:
            chunk["page"] = page_num
        result.extend(chunks)
    return result
